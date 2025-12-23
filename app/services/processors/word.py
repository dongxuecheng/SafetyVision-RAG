"""Word document processors"""

import subprocess
import shutil
import tempfile
import os
from pathlib import Path
from typing import List, Dict, Any

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

from .base import DocumentProcessor


class WordProcessor(DocumentProcessor):
    """Word (.docx) 处理：传统Chunking策略"""

    @staticmethod
    def process(
        file_path: str,
        metadata: Dict[str, Any],
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[Document]:
        """
        Process DOCX file using traditional chunking strategy

        Args:
            file_path: Path to DOCX file
            metadata: Metadata to attach to documents
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks

        Returns:
            List of Document objects with chunked content
        """
        doc = DocxDocument(file_path)
        content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        document = Document(page_content=content, metadata=metadata)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "],
        )
        return splitter.split_documents([document])


class LegacyWordProcessor(DocumentProcessor):
    """旧版Word (.doc) 处理：优先使用LibreOffice，fallback到antiword"""

    @staticmethod
    def _try_libreoffice_convert(file_path: str) -> str:
        """
        使用 LibreOffice 将 .doc 转换为 .docx，然后提取文本

        LibreOffice 优势：
        - 支持 Visio 嵌入图、OLE 对象
        - 支持横向页面和复杂布局
        - 兼容性最强

        Args:
            file_path: .doc 文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 转换失败
        """
        # 查找 LibreOffice 可执行文件
        libreoffice_paths = [
            "libreoffice",
            "soffice",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/local/bin/libreoffice",
            "/usr/local/bin/soffice",
            "/opt/libreoffice/program/soffice",
        ]

        soffice_path = None
        for path in libreoffice_paths:
            if shutil.which(path):
                soffice_path = path
                break

        if not soffice_path:
            raise ValueError("LibreOffice 未安装。请安装: apt-get install libreoffice")

        # 创建临时目录用于转换
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)

            try:
                # 使用 LibreOffice 无头模式转换 .doc -> .docx
                # --headless: 无GUI模式
                # --convert-to docx: 转换为docx格式
                # --outdir: 输出目录
                result = subprocess.run(
                    [
                        soffice_path,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        str(temp_dir_path),
                        file_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,  # 2分钟超时
                    check=False,
                )

                if result.returncode != 0:
                    error_msg = (
                        result.stderr.strip()
                        if result.stderr
                        else result.stdout.strip()
                    )
                    raise ValueError(
                        f"LibreOffice 转换失败 (返回码: {result.returncode}): {error_msg}"
                    )

                # 查找转换后的 .docx 文件
                docx_files = list(temp_dir_path.glob("*.docx"))
                if not docx_files:
                    raise ValueError("LibreOffice 转换成功但未找到输出文件")

                converted_docx = docx_files[0]

                # 使用 python-docx 提取文本
                doc = DocxDocument(str(converted_docx))
                content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

                if not content or not content.strip():
                    raise ValueError("转换后的文档内容为空")

                return content

            except subprocess.TimeoutExpired:
                raise ValueError("LibreOffice 转换超时（超过120秒）")
            except Exception as e:
                if isinstance(e, ValueError):
                    raise
                raise ValueError(f"LibreOffice 处理出错: {str(e)}")

    @staticmethod
    def _try_antiword(file_path: str) -> str:
        """
        使用 antiword 提取 .doc 文本（fallback方案）

        antiword 限制：
        - 不支持 Visio 图、OLE 对象
        - 不支持横向页面
        - 只支持简单的 Word 97-2003 格式

        Args:
            file_path: .doc 文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 提取失败
        """
        antiword_path = shutil.which("antiword")
        if not antiword_path:
            raise ValueError("antiword 工具未安装。请安装: apt-get install antiword")

        try:
            result = subprocess.run(
                [antiword_path, file_path],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )

            if result.returncode != 0:
                error_msg = result.stderr.strip() if result.stderr else "未知错误"
                raise ValueError(
                    f"antiword 执行失败 (返回码: {result.returncode}): {error_msg}"
                )

            content = result.stdout

            if not content or not content.strip():
                raise ValueError("antiword 无法提取文本，文件可能为空或格式不支持")

            return content

        except subprocess.TimeoutExpired:
            raise ValueError("antiword 处理超时（超过60秒）")
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"antiword 处理出错: {str(e)}")

    @staticmethod
    def process(
        file_path: str,
        metadata: Dict[str, Any],
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[Document]:
        """
        Process legacy DOC file with fallback strategy

        策略：
        1. 优先使用 LibreOffice（支持复杂格式）
        2. LibreOffice 失败时 fallback 到 antiword（仅支持简单格式）

        Args:
            file_path: Path to DOC file
            metadata: Metadata to attach to documents
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks

        Returns:
            List of Document objects with chunked content

        Raises:
            ValueError: If all conversion methods fail
        """
        content = None
        errors = []

        # 策略1: 尝试 LibreOffice（推荐，支持复杂文档）
        try:
            content = LegacyWordProcessor._try_libreoffice_convert(file_path)
            metadata["conversion_method"] = "libreoffice"
        except ValueError as e:
            errors.append(f"LibreOffice: {str(e)}")

        # 策略2: LibreOffice 失败，尝试 antiword（fallback）
        if not content:
            try:
                content = LegacyWordProcessor._try_antiword(file_path)
                metadata["conversion_method"] = "antiword"
            except ValueError as e:
                errors.append(f"antiword: {str(e)}")

        # 所有方法都失败
        if not content:
            error_summary = "\n".join(f"  - {err}" for err in errors)
            raise ValueError(
                f"处理 .doc 文件失败，已尝试所有方法:\n{error_summary}\n\n"
                f"建议：\n"
                f"1. 如果文档包含 Visio图/复杂对象，请安装 LibreOffice\n"
                f"2. 或将 .doc 文件手动转换为 .docx 格式后重新上传"
            )

        document = Document(page_content=content, metadata=metadata)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "],
        )
        return splitter.split_documents([document])
