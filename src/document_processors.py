"""Document processors for multi-format support"""

from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import openpyxl
from docx import Document as DocxDocument
import xlrd
import subprocess


class DocumentProcessorFactory:
    """Factory for creating document processors based on file type"""

    @staticmethod
    def process(
        file_path: str,
        metadata: Dict[str, Any],
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[Document]:
        """Process document based on file extension"""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return PDFProcessor.process(file_path, metadata, chunk_size, chunk_overlap)
        elif ext == ".docx":
            return WordProcessor.process(file_path, metadata, chunk_size, chunk_overlap)
        elif ext == ".doc":
            return LegacyWordProcessor.process(
                file_path, metadata, chunk_size, chunk_overlap
            )
        elif ext == ".xlsx":
            return ExcelProcessor.process(file_path, metadata)
        elif ext == ".xls":
            return LegacyExcelProcessor.process(file_path, metadata)
        else:
            raise ValueError(
                f"不支持的文件类型: {ext}。支持的格式: .pdf, .docx, .doc, .xlsx, .xls"
            )


class PDFProcessor:
    """PDF处理：传统Chunking策略"""

    @staticmethod
    def process(
        file_path: str, metadata: Dict[str, Any], chunk_size: int, chunk_overlap: int
    ) -> List[Document]:
        loader = PyPDFLoader(file_path)
        docs = loader.load()

        for doc in docs:
            doc.metadata.update(metadata)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "],
        )
        return splitter.split_documents(docs)


class WordProcessor:
    """Word处理：传统Chunking策略"""

    @staticmethod
    def process(
        file_path: str, metadata: Dict[str, Any], chunk_size: int, chunk_overlap: int
    ) -> List[Document]:
        doc = DocxDocument(file_path)
        content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        document = Document(page_content=content, metadata=metadata)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "],
        )
        return splitter.split_documents([document])


class ExcelProcessor:
    """
    Excel处理：Row-to-Text策略
    将每行数据结合表头转换为语义文本，保留结构信息
    """

    @staticmethod
    def process(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        documents = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))

            if not rows:
                continue

            # 第一行作为表头
            headers = [str(h).strip() if h else "" for h in rows[0]]

            # 处理数据行
            for row_idx, row in enumerate(rows[1:], start=2):
                row_data = []
                for header, value in zip(headers, row):
                    if header and value is not None:
                        value_str = (
                            value.strftime("%Y-%m-%d %H:%M:%S")
                            if isinstance(value, datetime)
                            else str(value).strip()
                        )
                        if value_str:
                            row_data.append(f"{header}为{value_str}")

                if row_data:
                    # 生成语义文本
                    text = f"第{row_idx}行数据：{'，'.join(row_data)}。"
                    doc_metadata = {
                        **metadata,
                        "sheet_name": sheet_name,
                        "row_number": row_idx,
                        "content_type": "excel_row",
                    }
                    documents.append(Document(page_content=text, metadata=doc_metadata))

        wb.close()
        return documents


class LegacyWordProcessor:
    """旧版Word(.doc)处理：使用antiword转换为文本，传统Chunking策略"""

    @staticmethod
    def process(
        file_path: str, metadata: Dict[str, Any], chunk_size: int, chunk_overlap: int
    ) -> List[Document]:
        import shutil

        # 检查 antiword 是否安装
        antiword_path = shutil.which("antiword")
        if not antiword_path:
            raise ValueError(
                "antiword 工具未安装或未找到。"
                "请确保系统已安装 antiword: apt-get install antiword"
            )

        try:
            # 使用 antiword 命令行工具提取文本
            result = subprocess.run(
                [antiword_path, file_path],
                capture_output=True,
                text=True,
                timeout=60,  # 设置超时保护
                check=False,  # 不自动抛出异常，手动处理
            )

            # 检查返回码
            if result.returncode != 0:
                error_msg = result.stderr.strip() if result.stderr else "未知错误"
                raise ValueError(
                    f"antiword 执行失败 (返回码: {result.returncode}): {error_msg}"
                )

            content = result.stdout

        except subprocess.TimeoutExpired:
            raise ValueError("处理 .doc 文件超时，文件可能太大或格式复杂")
        except Exception as e:
            raise ValueError(f"处理 .doc 文件时出错: {str(e)}")

        if not content or not content.strip():
            raise ValueError("无法从 .doc 文件中提取文本内容，文件可能为空或格式不支持")

        document = Document(page_content=content, metadata=metadata)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "],
        )
        return splitter.split_documents([document])


class LegacyExcelProcessor:
    """旧版Excel(.xls)处理：Row-to-Text策略"""

    @staticmethod
    def process(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        workbook = xlrd.open_workbook(file_path)
        documents = []

        for sheet in workbook.sheets():
            sheet_name = sheet.name

            if sheet.nrows == 0:
                continue

            # 第一行作为表头
            headers = [
                (
                    str(sheet.cell_value(0, col)).strip()
                    if sheet.cell_value(0, col)
                    else ""
                )
                for col in range(sheet.ncols)
            ]

            # 处理数据行
            for row_idx in range(1, sheet.nrows):
                row_data = []
                for col_idx, header in enumerate(headers):
                    if header and col_idx < sheet.ncols:
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value:
                            # 处理日期类型
                            if sheet.cell_type(row_idx, col_idx) == xlrd.XL_CELL_DATE:
                                date_tuple = xlrd.xldate_as_tuple(
                                    cell_value, workbook.datemode
                                )
                                value_str = f"{date_tuple[0]}-{date_tuple[1]:02d}-{date_tuple[2]:02d}"
                            else:
                                value_str = str(cell_value).strip()

                            if value_str:
                                row_data.append(f"{header}为{value_str}")

                if row_data:
                    # 生成语义文本
                    text = f"第{row_idx + 1}行数据：{'，'.join(row_data)}。"
                    doc_metadata = {
                        **metadata,
                        "sheet_name": sheet_name,
                        "row_number": row_idx + 1,
                        "content_type": "excel_row",
                    }
                    documents.append(Document(page_content=text, metadata=doc_metadata))

        return documents
